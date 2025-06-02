"""
GUI Word Document Generator

This script reads user data from an Excel sheet and fills a Word template for each entry.
The user selects the Excel file, Word template, and output folder through a graphical interface.
"""

import os
import datetime
import openpyxl
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

# Step 1: Read data from Excel file
def read_excel_data(excel_file_path):
    wb = openpyxl.load_workbook(excel_file_path)
    ws = wb.active
    data = []

    headers = [cell.value for cell in ws[1]]
    data.append(headers)

    for row in ws.iter_rows(min_row=2, values_only=True):
        formatted_row = []
        for cell in row:
            if isinstance(cell, (datetime.datetime, datetime.date)):
                formatted_row.append(cell.strftime('%Y-%m-%d'))
            else:
                formatted_row.append(str(cell))
        data.append(formatted_row)

    return data

# Step 2: Replace placeholders in Word template
def replace_placeholders(template_path, output_path, replacements):
    doc = Document(template_path)

    def replace_text_in_paragraph(paragraph, replacements):
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

    def replace_text_in_table(table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        replace_text_in_table(table, replacements)

    doc.save(output_path)

# Step 3: Generate documents
def fill_word_template_from_excel(excel_file_path, template_doc_path, output_dir):
    data = read_excel_data(excel_file_path)
    headers = data[0]
    rows = data[1:]

    for row in rows:
        replacements = {
            "{{Name and Surname}}": row[0],
            "{{Passport}}": row[3],
            "{{DoB}}": row[1],
            "{{Citizenship}}": row[2],
            "{{Uni}}": row[4]
        }

        output_file_name = f"{row[0].replace(' ', '_')}.docx"
        output_file_path = os.path.join(output_dir, output_file_name)
        replace_placeholders(template_doc_path, output_file_path, replacements)
        print(f"Created document: {output_file_name}")
    messagebox.showinfo("Done", "Documents generated successfully!")

# GUI
def run_gui():
    root = tk.Tk()
    root.title("Word Document Generator")
    root.geometry("600x400")

    excel_path = tk.StringVar()
    word_template_path = tk.StringVar()
    output_folder_path = tk.StringVar()

    def browse_excel():
        path = filedialog.askopenfilename(title="Select Excel File", filetypes=[("Excel Files", "*.xlsx")])
        excel_path.set(path)

    def browse_template():
        path = filedialog.askopenfilename(title="Select Word Template", filetypes=[("Word Files", "*.docx")])
        word_template_path.set(path)

    def browse_output():
        path = filedialog.askdirectory(title="Select Output Folder")
        output_folder_path.set(path)

    def generate_docs():
        if not excel_path.get() or not word_template_path.get() or not output_folder_path.get():
            messagebox.showerror("Missing Info", "Please select all paths.")
            return

        fill_word_template_from_excel(excel_path.get(), word_template_path.get(), output_folder_path.get())

    # GUI layout
    tk.Label(root, text="Excel File:").pack(pady=5)
    tk.Entry(root, textvariable=excel_path, width=50).pack()
    tk.Button(root, text="Browse", command=browse_excel).pack(pady=2)

    tk.Label(root, text="Word Template:").pack(pady=5)
    tk.Entry(root, textvariable=word_template_path, width=50).pack()
    tk.Button(root, text="Browse", command=browse_template).pack(pady=2)

    tk.Label(root, text="Output Folder:").pack(pady=5)
    tk.Entry(root, textvariable=output_folder_path, width=50).pack()
    tk.Button(root, text="Browse", command=browse_output).pack(pady=2)

    tk.Button(root, text="Generate Documents", command=generate_docs, bg="green", fg="white").pack(pady=15)

    root.mainloop()

if __name__ == "__main__":
    run_gui()
